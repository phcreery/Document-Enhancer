#Document Enhancer

import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import random
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
import re

import os
import sys
import torch
import random
import argparse
import numpy as np
from gpt2Pytorch.GPT2.model import (GPT2LMHeadModel)
from gpt2Pytorch.GPT2.utils import load_weight
from gpt2Pytorch.GPT2.config import GPT2Config
from gpt2Pytorch.GPT2.sample import sample_sequence
from gpt2Pytorch.GPT2.encoder import get_encoder

from gpt2Pytorch.mainLib import *

paragraph = 0
word = 0
titleLine=0

document = Document()
doc = docx.Document('test.docx')

def formatStyles():
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0,0,0)
    font.underline = False
    style.paragraph_format.line_spacing = 2

    styles = document.styles
    styles['Title'].delete()
    style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)

    style = document.styles['Title']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(20)
    font.color.rgb = RGBColor(0,0,0)
    font.underline = False
    style.paragraph_format.line_spacing = 2

def addheader():
    section = document.sections[0]
    header = section.header
    head = header.paragraphs[0]
    list = readParagraph(0)
    #document.add_paragraph("Testing", )
    #name = list[0]
    head.text = list.split( )[1]
    #head.text = readWord(0,1)
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def addheading():
    list = readHeading()
    #print(list)
    #headingDate()
    #for i in list:
    #    heading = document.add_paragraph(i)
    heading = document.add_paragraph(list[0])
    heading = document.add_paragraph(list[1])
    heading = document.add_paragraph(list[2])
    heading = document.add_paragraph(headingDate(list[3]))

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.line_spacing = 2


def addtitle():
    #document.add_heading('Document Title', 0)
    title = document.add_paragraph(readParagraph(titleLine), style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

#Adds body paragraphs into the body with new format.
def addbody():
    list = readBody()
    for i in list:
        #print(list[i])
        paragraph = document.add_paragraph('\t' + i[:-2])
    for paragraph_text in AIconverter(readBody()[0]).split('\n\n'):
        #print(paragraph_text.strip())
        paragraph = document.add_paragraph("\t"+paragraph_text.strip())
        #body.append(paragraph(paragraph_text.strip()))

    #paragraph = document.add_paragraph(AItext)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 2


#adds new page and titles it "Works Cited"
def addReferences():
    document.add_page_break()
    wstitle = document.add_paragraph('Works Cited', style='Title')
    wstitle.alignment = WD_ALIGN_PARAGRAPH.CENTER


#returns heading as list
def readHeading():
    global titleLine
    #doc = docx.Document('test.docx')
    #headingList = [doc.paragraphs[0].text, doc.paragraphs[1].text, doc.paragraphs[2].text, doc.paragraphs[3].text]
    headingList=[]
    i = 0
    while i <= 4:
        #print(doc.paragraphs[i].alignment)
        if doc.paragraphs[i].alignment != WD_ALIGN_PARAGRAPH.CENTER:
            #print(str(i)+" line is not title")
            headingList.append(doc.paragraphs[i].text)
        else:
            if i < 4:
                headingList = ["FIRST LAST", "(PROF,Mr.,Mrs,Ms) NAME", "CLASS", "DD MMM YYYY"]
            titleLine = i
            break
        i=i+1
    print(headingList)
    return headingList;

#returns any paragraph by its number
def readParagraph(paragraph):
    #doc = docx.Document('test.docx')
    return doc.paragraphs[paragraph].text

#returns list of all paragraphs after Title
def readBody():
    #doc = docx.Document('test.docx')
    list = []
    i=titleLine+1
    print(i)
    while i < len(doc.paragraphs):
        list.append(doc.paragraphs[i].text)
        print(list)
        i+=1
    return list


def headingDate(date):
    #print(date)
    #date="mar 2020 24"
    try:
        day = re.search("([^\d])([0-2]|[0-2][0-9])([^\d])" , " "+date+" ")
        year = re.search("[2-9][0-9][0-9][0-9]" , date)
        month = re.search("[^\s\d][^\s\d][^\s\d]" , date)
        newdate=day.group()[1:-1]+" "+month.group().capitalize()+". "+year.group()
        #print(newdate)
    except:
        print("error, unable to corrct date ... ignoring")
        newdate=date
    return newdate


#returns true if a word is ignorable and false if important
def ignorable(word):
    ignore = ["The", "the", "To", "to", "Of", "of", "Be", "be", "and", "A", "a", "That", "that", "Have", "have", "I",
              "It", "it", "For", "for", "Not", "not", "With", "with", "You", "you", "As", "as", "Do", "do", "At", "at"
              "This", "this", "By", "by", "or", "An", "an", "From", "from", "Will", "will", "Is", "is"]
    for x in range( 0, len(ignore) ):
        if(word == ignore[x]):
            return True
    return False

#takes in a string and returns a dictionary on the word count of each word
def getRepetitive( text ):
    unique = {}

    for word in text:
        if ignorable(word) == False:
            if len(unique) == 0:
                unique[word] = 1
            else:
                for y in list(unique):
                    if(word == y):
                        unique[word] += 1
                        break
                    unique[word] = 1
    return unique

# Hunter Testing

def phrases(string):
    words = string.split()
    result = []
    for number in range(len(words)):
        for start in range(len(words)-number):
             result.append(" ".join(words[start:start+number+1]))
    return result




#MAIN
def main():
    formatStyles()
    #addheader()
    addheading()
    addtitle()
    #readBody()      #just testing
    addbody()
    addReferences()
    document.save('mla.docx')

main()
